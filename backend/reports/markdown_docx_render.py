"""
Renders parsed markdown blocks (markdown_blocks.py) into python-docx
paragraphs/lists/tables — counterpart to markdown_pdf_render.py for the
Word export. Same underlying problem: add_paragraph(response) collapsed
the response's real paragraph/heading/list/table structure (python-docx
paragraphs don't turn embedded "\n" into visual line breaks either) and
left literal "##"/"**"/"|" markdown syntax in the text.

Uses Word's built-in "List Bullet"/"List Number" paragraph styles for real
list rendering rather than hand-building numbering.xml — every default
python-docx Document() template already ships these, no custom numbering
definitions needed.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from backend.reports.markdown_blocks import parse_inline_runs, parse_markdown_blocks

# Spacing kept tight for the same reason as markdown_pdf_render.py: a real
# analyst response bucket runs to dozens of blocks per response, and a few
# points of paragraph spacing compounds fast across hundreds of responses.
_HEADING_SIZES = {1: 12, 2: 11, 3: 10.5, 4: 10, 5: 10, 6: 10}


def _add_runs(paragraph, runs, size_pt, color, bold_all=False):
    for text, is_bold in runs:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.font.bold = bold_all or is_bold


def _add_hr(doc, color_hex="D1D5DB"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def render_markdown_to_docx(doc, text, body_color, heading_color,
                            table_header_bg="2D5A8E"):
    blocks = parse_markdown_blocks(text)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            size = _HEADING_SIZES.get(block["level"], 10)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            _add_runs(p, block["runs"], size, heading_color, bold_all=True)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_runs(p, block["runs"], 9, body_color)

        elif btype == "hr":
            _add_hr(doc)

        elif btype in ("bullet_list", "numbered_list"):
            style = "List Bullet" if btype == "bullet_list" else "List Number"
            for item_runs in block["items"]:
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(0)
                _add_runs(p, item_runs, 9, body_color)

        elif btype == "table":
            header = block["header"]
            rows = block["rows"]
            if not header:
                continue
            n_cols = len(header)
            tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
            tbl.style = "Table Grid"

            for c, cell_text in enumerate(header):
                cell = tbl.rows[0].cells[c]
                _add_runs(cell.paragraphs[0], parse_inline_runs(cell_text),
                          8.5, RGBColor(0xFF, 0xFF, 0xFF), bold_all=True)
                _shade_cell(cell, table_header_bg)

            for r, row in enumerate(rows, start=1):
                padded = (row + [""] * n_cols)[:n_cols]
                for c, cell_text in enumerate(padded):
                    cell = tbl.rows[r].cells[c]
                    _add_runs(cell.paragraphs[0], parse_inline_runs(cell_text), 8.5, body_color)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)
