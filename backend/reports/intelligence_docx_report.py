"""
Word (.docx) report generator for Atlas AI Intelligence Engine reports.
Uses python-docx for clean, editable output.

Data shapes match intelligence_repository selects:
    run         — (run_id, provider, model, target_brand, started_at, …, duration_seconds)
    briefing    — (product_summary, persona_summary, journey_summary,
                   opportunities_text, executive_briefing, created_at)
    results     — [(analyst_name, prompt, response, collected_at), …]
    opportunities — [(id, title, evidence, description, status), …]
"""
from __future__ import annotations

from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from backend.intelligence.opportunity_ranking import rank_opportunities
from backend.reports.provenance import (
    SECTION_INTROS, UNVERIFIED_BADGE, UNVERIFIED_NOTE,
)
from backend.reports.briefing_sections import split_briefing_sections
from backend.reports.markdown_docx_render import render_markdown_to_docx


# ── Palette ────────────────────────────────────────────────────────────────────
_NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
_BLUE  = RGBColor(0x0B, 0x84, 0xFF)
_GRAY  = RGBColor(0x6B, 0x72, 0x80)
_DARK  = RGBColor(0x11, 0x18, 0x27)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN = RGBColor(0x16, 0xA3, 0x4A)
_AMBER = RGBColor(0xF5, 0x9E, 0x0B)

_NAVY_HEX  = "1E3A5F"
_BLUE_HEX  = "0B84FF"
_BG_HEX    = "F3F4F6"
_LIGHT_HEX = "E5E7EB"


class IntelligenceDocxReport:
    """
    Generate an editable .docx from an Intelligence Engine analysis run.

    Usage:
        rpt = IntelligenceDocxReport(
            run=run_tuple,
            briefing=briefing_tuple,
            results=results_list,
            opportunities=opp_list,
            target_brand="Firman",
        )
        rpt.generate("/path/to/report.docx")
    """

    def __init__(self, run, briefing, results, opportunities,
                 target_brand: str = "", full_export: bool = False):
        """full_export: see IntelligencePDFReport's docstring for this
        parameter — same meaning here."""
        self.run           = run or ()
        self.briefing      = briefing or ()
        self.results       = results or []
        self.opportunities = opportunities or []
        self.full_export   = full_export
        self.target_brand  = (
            target_brand
            or (run[3] if run and len(run) > 3 else "")
            or "Target Brand"
        )
        self.generated_at  = datetime.now()

    # ── Public ─────────────────────────────────────────────────────────────────

    def generate(self, output_path: str) -> None:
        doc = Document()
        self._set_page_layout(doc)
        self._set_styles(doc)

        self._write_cover(doc)
        self._write_briefing(doc)
        self._write_analyst_sections(doc)
        self._write_opportunities(doc)

        doc.save(output_path)

    # ── Document setup ─────────────────────────────────────────────────────────

    def _set_page_layout(self, doc: Document):
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    def _set_styles(self, doc: Document):
        styles = doc.styles

        normal = styles['Normal']
        normal.font.name  = 'Calibri'
        normal.font.size  = Pt(10)
        normal.font.color.rgb = _DARK

        for h_name, size, bold in [
            ('Heading 1', 16, True),
            ('Heading 2', 13, True),
            ('Heading 3', 11, True),
        ]:
            style = styles[h_name]
            style.font.name  = 'Calibri'
            style.font.size  = Pt(size)
            style.font.bold  = bold
            style.font.color.rgb = _NAVY
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after  = Pt(4)

    # ── Cover ──────────────────────────────────────────────────────────────────

    def _write_cover(self, doc: Document):
        run_tup = self.run
        tb      = self.target_brand

        # Top accent line (colored paragraph border)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after  = Pt(0)
        self._set_bottom_border(p, _BLUE_HEX, size=12)

        # "ATLAS AI" label
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run("ATLAS AI")
        run.font.name   = 'Calibri'
        run.font.size   = Pt(10)
        run.font.bold   = True
        run.font.color.rgb = _BLUE

        # Report title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(12)
        run = p.add_run("Intelligence Engine Report")
        run.font.name = 'Calibri'
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = _DARK

        # Target brand
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run("Target Brand")
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = _GRAY

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(16)
        run = p.add_run(tb)
        run.font.name = 'Calibri'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = _NAVY

        # Divider
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(12)
        self._set_bottom_border(p, _LIGHT_HEX)

        # Metadata table
        provider = run_tup[1].capitalize() if run_tup and len(run_tup) > 1 else "—"
        model    = run_tup[2] if run_tup and len(run_tup) > 2 else "—"
        started  = (run_tup[4][:16].replace("T", "  ")
                    if run_tup and len(run_tup) > 4 and run_tup[4] else "—")
        duration = (f"{run_tup[7]:.0f}s"
                    if run_tup and len(run_tup) > 7 and run_tup[7] else "—")

        meta_rows = [
            ("Generated",           self.generated_at.strftime("%B %d, %Y  ·  %H:%M")),
            ("Run Date",            started),
            ("AI Provider",         provider),
            ("Model",               model),
            ("Duration",            duration),
            ("Analyst Responses",   str(len(self.results))),
            ("Opportunities Found", str(len(self.opportunities))),
        ]
        tbl = doc.add_table(rows=len(meta_rows), cols=2)
        tbl.style = 'Table Grid'
        for i, (label, value) in enumerate(meta_rows):
            row = tbl.rows[i]
            lc = row.cells[0]
            vc = row.cells[1]

            lc.width = Inches(1.8)
            vc.width = Inches(4.7)

            lc_p = lc.paragraphs[0]
            lc_r = lc_p.add_run(label)
            lc_r.font.name = 'Calibri'
            lc_r.font.size = Pt(9)
            lc_r.font.bold = True
            lc_r.font.color.rgb = _GRAY

            vc_p = vc.paragraphs[0]
            vc_r = vc_p.add_run(value)
            vc_r.font.name = 'Calibri'
            vc_r.font.size = Pt(9)
            vc_r.font.color.rgb = _DARK

            self._shade_cell(lc, "FAFAFA")
            self._shade_cell(vc, "FFFFFF")

        doc.add_page_break()

    # ── Executive Briefing ─────────────────────────────────────────────────────

    def _write_briefing(self, doc: Document):
        txt = self.briefing[4] if self.briefing and len(self.briefing) > 4 else ""
        if not txt:
            return

        doc.add_heading("Executive Briefing", level=1)

        # The briefing prompt deliberately forbids markdown and instead uses
        # plain "SECTION NAME\nBody text" sections (briefing_sections.py) —
        # give each section's header real visual distinction instead of one
        # dense, undifferentiated wall of paragraphs.
        for header, body in split_briefing_sections(txt):
            if header:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(header)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = _NAVY
            if body:
                p = doc.add_paragraph(body)
                p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # ── Analyst Q&A sections ───────────────────────────────────────────────────

    _MAX_QA_PAIRS_SHOWN = 5

    def _write_analyst_sections(self, doc: Document):
        # Shared with the PDF via provenance.py so both describe these blocks
        # as CAPTURED EVIDENCE rather than Atlas analysis (R7).
        _SECTIONS = [
            (name, SECTION_INTROS[name])
            for name in ("Product Intelligence", "Consumer Personas", "Buying Journey")
        ]

        by_analyst: dict[str, list[tuple[str, str]]] = {}
        for analyst_name, prompt, response, _ in self.results:
            by_analyst.setdefault(analyst_name, []).append((prompt, response))

        for section_name, intro in _SECTIONS:
            all_pairs = by_analyst.get(section_name, [])
            if not all_pairs:
                continue
            # Same "too much raw detail for a printed report" cap as the PDF
            # export and Strategic Opportunities below — the Executive
            # Briefing is the actual synthesis of this data; these are
            # representative supporting examples, not primary reading.
            qa_cap = len(all_pairs) if self.full_export else self._MAX_QA_PAIRS_SHOWN
            pairs = all_pairs[:qa_cap]

            doc.add_heading(section_name, level=1)
            intro_text = intro
            if len(all_pairs) > qa_cap:
                intro_text += (
                    f" Showing {len(pairs)} of {len(all_pairs)} representative responses "
                    "analyzed for this section; see the Product/Personas/Journey tabs in "
                    "the app for the complete set."
                )
            p = doc.add_paragraph(intro_text)
            p.paragraph_format.space_after = Pt(4)

            # R7: provenance banner between the intro and the first verbatim
            # answer — bold amber so it cannot be mistaken for body copy or
            # skimmed past on the way into unverified model output.
            warn = doc.add_paragraph()
            badge = warn.add_run(f"{UNVERIFIED_BADGE} — ")
            badge.bold = True
            badge.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            note = warn.add_run(UNVERIFIED_NOTE)
            note.font.color.rgb = RGBColor(0x7C, 0x4A, 0x03)
            for run in (badge, note):
                run.font.size = Pt(8.5)
            warn.paragraph_format.space_after = Pt(10)

            for prompt, response in pairs:
                # Prompt line
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(prompt)
                run.font.bold  = True
                run.font.size  = Pt(9)
                run.font.color.rgb = _NAVY

                # Response: real markdown (headers/bold/bullets/tables)
                # rendered as actual formatting instead of one add_paragraph()
                # call, which showed literal ##/**/| syntax AND collapsed the
                # response's real structure (docx paragraphs don't turn
                # embedded newlines into visual line breaks either).
                if response:
                    render_markdown_to_docx(doc, response, _DARK, _NAVY)
                else:
                    doc.add_paragraph("(no response)")

            doc.add_page_break()

    # ── Strategic Opportunities ────────────────────────────────────────────────

    _MAX_OPPORTUNITIES_SHOWN = 10

    def _write_opportunities(self, doc: Document):
        # Ranked (evidence-count-backed findings first), not just parse
        # order -- see opportunity_ranking.py. Capped for the same reason
        # as the PDF export.
        all_opps = rank_opportunities(self.opportunities)
        opp_cap = len(all_opps) if self.full_export else self._MAX_OPPORTUNITIES_SHOWN
        opps = all_opps[:opp_cap]
        if not opps:
            return

        doc.add_heading("Strategic Opportunities", level=1)
        intro = (
            f"The following opportunities were identified through AI analysis of how "
            f"{self.target_brand} and competitors are described across consumer research "
            "scenarios. Each represents an area where improved content, positioning, or "
            "product strategy could increase AI visibility. Ranked by strength of evidence "
            "— findings citing a specific count (e.g. \"0 of 84 responses\") are shown first."
        )
        if len(all_opps) > opp_cap:
            intro += (
                f" Showing the top {opp_cap} of {len(all_opps)} "
                "opportunities identified from this run; see the Opportunities tab in "
                "the app for the complete list."
            )
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(12)

        _STATUS_LABELS = {
            "new":         "New",
            "in_progress": "In Progress",
            "done":        "Done",
        }

        for idx, row in enumerate(opps, 1):
            # Sliced, not unpacked to a fixed arity — opps can come from either
            # get_opportunities_for_run() (5 columns) or get_all_opportunities()
            # (6, trailing created_date); see intelligence_pdf_report.py's
            # matching fix for the real crash this caused on Full export.
            opp_id, title, evidence, description, status = row[:5]
            # Opportunity number + title
            doc.add_heading(f"{idx}.  {title}", level=2)

            status_label = _STATUS_LABELS.get(status or "new", "New")

            rows = []
            if status:
                rows.append(("Status", status_label))
            if evidence:
                rows.append(("Evidence", evidence))
            if description:
                # "Action" — matches the on-screen label (intelligence_page.py's
                # opportunity cards) for this same field; previously said
                # "Description" here, which didn't match the screen.
                rows.append(("Action", description))

            if rows:
                tbl = doc.add_table(rows=len(rows), cols=2)
                tbl.style = 'Table Grid'
                for i, (label, value) in enumerate(rows):
                    row = tbl.rows[i]
                    lc = row.cells[0]
                    vc = row.cells[1]

                    lc.width = Inches(1.2)
                    vc.width = Inches(5.3)

                    lc_p = lc.paragraphs[0]
                    lc_r = lc_p.add_run(label)
                    lc_r.font.bold = True
                    lc_r.font.size = Pt(8.5)
                    lc_r.font.color.rgb = _GRAY

                    vc_p = vc.paragraphs[0]
                    vc_r = vc_p.add_run(value)
                    vc_r.font.size = Pt(9)
                    vc_r.font.color.rgb = _DARK

                    self._shade_cell(lc, "F9FAFB")
                    self._shade_cell(vc, "FFFFFF")

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(12)

    # ── OOXML helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def _set_bottom_border(paragraph, color_hex: str, size: int = 6):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _shade_cell(cell, fill_hex: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcPr.append(shd)
