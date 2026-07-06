"""
Tests for IntelligenceDocxReport (#33) — the Word export had zero coverage
despite being directly user-facing (Intelligence page's "Export Word"
button). Tuple shapes match the module's own documented format (run,
briefing, results, opportunities) rather than guessing.
"""
from docx import Document

from backend.reports.intelligence_docx_report import IntelligenceDocxReport


def _all_paragraph_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _heading_texts(doc: Document) -> list:
    return [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]


_RUN = ("run-1", "openai", "gpt-4.1-mini", "Firman", "2026-07-01T10:00:00", None, None, 26.4)
_BRIEFING = ("product summary", "persona summary", "journey summary",
             "opportunities text", "This is the executive briefing.\n\nSecond paragraph.",
             "2026-07-01T10:00:26")
_RESULTS = [
    ("Product Intelligence", "What features matter?", "Dual fuel and inverter tech.", "2026-07-01T10:00:05"),
    ("Consumer Personas", "Who buys generators?", "Homeowners and contractors.", "2026-07-01T10:00:10"),
]
_OPPORTUNITIES = [
    (1, "Improve Amazon presence", "Firman has fewer reviews than Honda", "Increase review volume", "new"),
    (2, "Fix spec accuracy", "AI cites wrong running watts", "Update product feed", "in_progress"),
]


def test_generate_produces_valid_docx_with_all_sections(tmp_path):
    out = tmp_path / "report.docx"
    IntelligenceDocxReport(
        run=_RUN, briefing=_BRIEFING, results=_RESULTS,
        opportunities=_OPPORTUNITIES, target_brand="Firman",
    ).generate(str(out))

    assert out.exists() and out.stat().st_size > 0
    doc = Document(str(out))
    headings = _heading_texts(doc)
    assert "Executive Briefing" in headings
    assert "Product Intelligence" in headings
    assert "Consumer Personas" in headings
    assert "Strategic Opportunities" in headings


def test_cover_page_shows_run_metadata(tmp_path):
    out = tmp_path / "cover.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, _RESULTS, _OPPORTUNITIES, "Firman").generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Firman" in text
    assert "Openai" in text  # provider.capitalize()
    assert "gpt-4.1-mini" in text
    assert "26s" in text  # duration formatted as f"{duration:.0f}s"
    assert "2" in text  # analyst response count / opportunities count


def test_analyst_result_with_unrecognized_name_is_silently_dropped(tmp_path):
    """Current, real behavior: _write_analyst_sections only recognizes
    exactly 'Product Intelligence'/'Consumer Personas'/'Buying Journey'
    (confirmed these match intelligence_service.py's actual bucket names).
    A result tagged with anything else — e.g. a typo, or a name from a
    future new analyst bucket that isn't wired into this report yet —
    is silently excluded rather than raising or showing an 'Other'
    section. Pinning this so it's a known, intentional limitation, not
    a silent data-loss surprise if it's ever hit in practice."""
    results = _RESULTS + [
        ("Brand Intelligence", "How is the brand perceived?", "Generally positive.", "2026-07-01T10:00:15"),
    ]
    out = tmp_path / "dropped.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, results, [], "Firman").generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Generally positive." not in text
    assert "Dual fuel and inverter tech." in text  # the recognized ones still appear


def test_opportunity_status_labels_render_correctly(tmp_path):
    out = tmp_path / "opps.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, [], _OPPORTUNITIES, "Firman").generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "New" in text
    assert "In Progress" in text
    assert "Improve Amazon presence" in text
    assert "Fix spec accuracy" in text


def test_full_export_disables_qa_pair_cap(tmp_path):
    results = [
        ("Product Intelligence", f"Question {i}?", f"Response body {i}.", "2026-07-01T10:00:00")
        for i in range(8)
    ]
    out = tmp_path / "full.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, results, [], "Firman", full_export=True).generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Showing" not in text
    for i in range(8):
        assert f"Question {i}?" in text


def test_full_export_disables_opportunities_cap(tmp_path):
    opps = [(i, f"Opp {i}", f"{i} of 20 responses", f"Action {i}", "new") for i in range(15)]
    out = tmp_path / "full_opps.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, [], opps, "Firman", full_export=True).generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Showing the top" not in text
    for i in range(15):
        assert f"Opp {i}" in text


def test_executive_briefing_sections_get_distinct_header_styling(tmp_path):
    briefing = (
        "product summary", "persona summary", "journey summary", "opportunities text",
        "VISIBILITY SNAPSHOT  \nFirman appeared in 12 of 84 responses.\n\n"
        "SENTIMENT  \nOut of 12 mentions, 4 were negative.",
        "2026-07-01T10:00:26",
    )
    out = tmp_path / "briefing_sections.docx"
    IntelligenceDocxReport(_RUN, briefing, [], [], "Firman").generate(str(out))
    doc = Document(str(out))
    paras = [p for p in doc.paragraphs if p.text.strip()]

    header_para = next(p for p in paras if p.text == "VISIBILITY SNAPSHOT")
    assert header_para.runs[0].font.bold is True
    body_para = next(p for p in paras if "Firman appeared in 12 of 84 responses." in p.text)
    assert body_para.text == "Firman appeared in 12 of 84 responses."


def test_analyst_section_caps_qa_pairs_shown_with_a_note(tmp_path):
    results = [
        ("Product Intelligence", f"Question {i}?", f"Response body {i}.", "2026-07-01T10:00:00")
        for i in range(8)
    ]
    out = tmp_path / "capped.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, results, [], "Firman").generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Showing 5 of 8" in text
    for i in range(5):
        assert f"Question {i}?" in text
    for i in range(5, 8):
        assert f"Question {i}?" not in text


def test_opportunities_section_is_ranked_and_capped(tmp_path):
    opps = [
        (1, "Qualitative gap", "Champion and Honda have thousands of reviews.", "Action A", "new"),
        (2, "Complete absence", "Firman is not mentioned (0 of 84 responses).", "Action B", "new"),
    ] + [
        (i, f"Filler {i}", "No count here.", f"Action {i}", "new") for i in range(3, 13)
    ]
    out = tmp_path / "ranked_opps.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, [], opps, "Firman").generate(str(out))
    doc = Document(str(out))
    text = _all_paragraph_text(doc)
    headings = _heading_texts(doc)

    assert "Showing the top 10 of 12" in text
    heading_order = [h for h in headings if "Complete absence" in h or "Qualitative gap" in h]
    assert "Complete absence" in heading_order[0]  # ranked first (0/84 ratio)


def test_opportunity_section_labels_field_action_not_description(tmp_path):
    """Regression: this field was labeled "Description" here but "Action" on
    the Intelligence page's opportunity cards — same underlying data, two
    different labels depending on where you looked."""
    out = tmp_path / "labels.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, [], _OPPORTUNITIES, "Firman").generate(str(out))
    text = _all_paragraph_text(Document(str(out)))

    assert "Action" in text
    assert "Description" not in text


def test_missing_briefing_text_skips_section_without_crashing(tmp_path):
    empty_briefing = ("", "", "", "", "", "")  # executive_briefing (index 4) blank
    out = tmp_path / "no_briefing.docx"
    IntelligenceDocxReport(_RUN, empty_briefing, _RESULTS, [], "Firman").generate(str(out))

    doc = Document(str(out))
    assert "Executive Briefing" not in _heading_texts(doc)


def test_empty_opportunities_skips_section_without_crashing(tmp_path):
    out = tmp_path / "no_opps.docx"
    IntelligenceDocxReport(_RUN, _BRIEFING, _RESULTS, [], "Firman").generate(str(out))

    doc = Document(str(out))
    assert "Strategic Opportunities" not in _heading_texts(doc)


def test_generate_does_not_crash_with_all_empty_inputs(tmp_path):
    """No run yet (e.g. report requested before any Intelligence Analysis
    has ever completed) — every field falls back to its documented
    default (run=(), briefing=(), results=[], opportunities=[])."""
    out = tmp_path / "empty.docx"
    IntelligenceDocxReport(run=None, briefing=None, results=None, opportunities=None).generate(str(out))
    assert out.exists() and out.stat().st_size > 0

    text = _all_paragraph_text(Document(str(out)))
    assert "Target Brand" in text  # falls back to the literal default label
